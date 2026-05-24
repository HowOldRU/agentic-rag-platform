"""
Word 文档提取工具 - 使用 python-docx 库
"""
import os
import uuid
import logging
from typing import Dict, Any
from abc import ABC
from docx import Document

logger = logging.getLogger(__name__)


class WordExtractionTool(ABC):
    """Word 提取工具，使用 python-docx 库解析 Word 文件"""
    name: str = "word_extractor"
    description: str = "Word 文件提取工具，支持提取 Word 中的文本和表格内容，支持 .docx 格式"

    def __init__(self):
        pass

    def is_word_file(self, file_name: str) -> bool:
        """检查是否为 Word 文件"""
        if not file_name:
            return False
        return file_name.lower().endswith((".docx", ".doc"))

    def extract_text_from_document(self, doc) -> str:
        """从 Word 文档中提取文本和表格内容"""
        content_parts = []

        # 提取段落文本
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content_parts.append(paragraph.text)

        # 提取表格内容
        for table in doc.tables:
            content_parts.append("\n[表格开始]")
            for row in table.rows:
                row_content = []
                for cell in row.cells:
                    row_content.append(cell.text.strip())
                content_parts.append("\t".join(row_content))
            content_parts.append("[表格结束]\n")

        return "\n".join(content_parts)

    def run(self, **kwargs) -> Dict[str, Any]:
        """执行 Word 提取，使用 python-docx 库"""
        file_content = kwargs.get("file_content")
        file_name = kwargs.get("file_name", "")
        temp_path = None

        try:
            # 验证文件类型
            if not self.is_word_file(file_name):
                return {
                    "success": False,
                    "error": f"不支持的文件类型，仅支持 DOCX 文件。当前文件: {file_name}"
                }

            # 检查文件内容是否为空
            if not file_content:
                return {
                    "success": False,
                    "error": "文件内容为空"
                }

            # 生成唯一的临时文件名以避免冲突
            file_ext = file_name.lower().split('.')[-1]
            temp_path = f"temp_{uuid.uuid4().hex}.{file_ext}"

            # 临时保存文件内容
            with open(temp_path, "wb") as f:
                f.write(file_content)

            # 检查文件是否创建成功且不为空
            if not os.path.exists(temp_path) or os.path.getsize(temp_path) == 0:
                return {
                    "success": False,
                    "error": "临时文件创建失败或文件为空"
                }

            # 使用 python-docx 库解析 Word 文档
            doc = Document(temp_path)

            # 提取文档内容
            doc_text = self.extract_text_from_document(doc)

            # 创建元数据
            metadata = {
                "file_type": "docx",
                "file_name": file_name,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }

            # 创建文档对象
            serialized_doc = {
                "page_content": doc_text,
                "metadata": metadata
            }

            result = {
                "success": True,
                "documents": [serialized_doc],
                "paragraph_count": len(doc.paragraphs)
            }

            return result

        except Exception as e:
            error_msg = f"Word 文件处理失败: {str(e)}"
            logger.error(error_msg)
            return {
                "success": False,
                "error": error_msg,
                "documents": []
            }
        finally:
            # 确保清理临时文件
            try:
                if temp_path and os.path.exists(temp_path):
                    os.remove(temp_path)
            except:
                pass
