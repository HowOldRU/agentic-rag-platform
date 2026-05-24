"""会话导出生成器 — 支持 PDF / DOCX / TXT 三种格式"""
import io
import re

from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.lib.colors import HexColor
from reportlab.platypus import Paragraph, Spacer, SimpleDocTemplate
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# reportlab 中文字体注册
try:
    pdfmetrics.registerFont(TTFont("SimHei", "SimHei.ttf"))
except Exception:
    # Windows 系统字体路径
    try:
        import os

        font_path = os.path.join(
            os.environ.get("SystemRoot", r"C:\Windows"),
            "Fonts",
            "simhei.ttf",
        )
        pdfmetrics.registerFont(TTFont("SimHei", font_path))
    except Exception:
        # 无法注册中文字体时 fallback 到默认
        pass

# 颜色常量
_USER_COLOR = HexColor("#1a73e8")  # 用户标签蓝色
_AI_COLOR = HexColor("#16a34a")  # AI 标签绿色
_TIME_COLOR = HexColor("#999999")  # 时间戳灰色
_USER_COLOR_RGB = RGBColor(0x1A, 0x73, 0xE8)
_AI_COLOR_RGB = RGBColor(0x16, 0xA3, 0x4A)
_TIME_COLOR_RGB = RGBColor(0x99, 0x99, 0x99)

# 用于清理 PPT 文件标记
_PPT_MARKER_RE = re.compile(r"<!--PPT_FILE:.+?-->")


def _clean_content(text: str) -> str:
    """清理消息内容中的特殊标记"""
    text = _PPT_MARKER_RE.sub("", text)
    return text.strip()


def _fmt_time(dt) -> str:
    """格式化时间戳"""
    return dt.strftime("%Y-%m-%d %H:%M")


def generate_txt(conv, messages) -> bytes:
    """生成 TXT 纯文本格式"""
    lines = [
        f"会话标题: {conv.title}",
        f"创建时间: {_fmt_time(conv.created_at)}",
        "",
    ]
    for msg in messages:
        role = "用户" if msg.role == "user" else "AI助手"
        lines.append(f"[{role}] {_fmt_time(msg.created_at)}")
        lines.append(_clean_content(msg.content))
        lines.append("")
    return "\n".join(lines).encode("utf-8")


def generate_pdf(conv, messages) -> bytes:
    """生成 PDF 格式（reportlab）"""
    buf = io.BytesIO()
    has_chinese_font = "SimHei" in pdfmetrics.getRegisteredFontNames()

    base_font = "SimHei" if has_chinese_font else "Helvetica"

    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )

    styles = {
        "title": ParagraphStyle(
            "Title",
            fontName=base_font,
            fontSize=16,
            leading=24,
            spaceAfter=4,
        ),
        "time": ParagraphStyle(
            "Time",
            fontName=base_font,
            fontSize=9,
            leading=14,
            textColor=_TIME_COLOR,
            spaceAfter=8,
        ),
        "user_label": ParagraphStyle(
            "UserLabel",
            fontName=base_font,
            fontSize=11,
            leading=16,
            textColor=_USER_COLOR,
            spaceBefore=10,
            spaceAfter=2,
        ),
        "ai_label": ParagraphStyle(
            "AILabel",
            fontName=base_font,
            fontSize=11,
            leading=16,
            textColor=_AI_COLOR,
            spaceBefore=10,
            spaceAfter=2,
        ),
        "body": ParagraphStyle(
            "Body",
            fontName=base_font,
            fontSize=10,
            leading=18,
            spaceAfter=6,
        ),
    }

    story = []
    story.append(Paragraph(_escape_xml(conv.title), styles["title"]))
    story.append(
        Paragraph(f"创建时间: {_fmt_time(conv.created_at)}", styles["time"])
    )

    for msg in messages:
        role = "用户" if msg.role == "user" else "AI助手"
        label_style = styles["user_label"] if msg.role == "user" else styles["ai_label"]
        story.append(
            Paragraph(
                f"{role} · {_fmt_time(msg.created_at)}",
                label_style,
            )
        )
        content = _clean_content(msg.content)
        if content:
            story.append(Paragraph(_escape_xml(content), styles["body"]))

    doc.build(story)
    return buf.getvalue()


def generate_docx(conv, messages) -> bytes:
    """生成 Word (DOCX) 格式（python-docx）"""
    doc = Document()

    # 标题
    title = doc.add_heading(conv.title, level=1)
    for run in title.runs:
        run.font.size = Pt(18)

    # 时间信息
    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"创建时间: {_fmt_time(conv.created_at)}")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = _TIME_COLOR_RGB

    for msg in messages:
        role = "用户" if msg.role == "user" else "AI助手"

        # 角色标签 + 时间戳
        label = doc.add_paragraph()
        label.space_before = Pt(12)
        role_run = label.add_run(f"[{role}]")
        role_run.bold = True
        role_run.font.size = Pt(11)
        role_run.font.color.rgb = (
            _USER_COLOR_RGB if msg.role == "user" else _AI_COLOR_RGB
        )
        time_run = label.add_run(f"  {_fmt_time(msg.created_at)}")
        time_run.font.size = Pt(9)
        time_run.font.color.rgb = _TIME_COLOR_RGB

        # 内容
        content = _clean_content(msg.content)
        if content:
            body = doc.add_paragraph(content)
            body.paragraph_format.space_after = Pt(4)
            for run in body.runs:
                run.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _escape_xml(text: str) -> str:
    """转义 XML 特殊字符（reportlab Paragraph 需要）"""
    text = text.replace("&", "&amp;")
    text = text.replace("<", "&lt;")
    text = text.replace(">", "&gt;")
    text = text.replace("\n", "<br/>")
    return text
